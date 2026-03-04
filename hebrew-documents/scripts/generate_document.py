#!/usr/bin/env python3
"""Generate Hebrew RTL Word documents with Israeli formatting.

Creates properly formatted Hebrew documents with RTL support,
Israeli date formats, and common document templates.

Usage:
    python generate_document.py --template letter --output output.docx
    python generate_document.py --template contract --output contract.docx
    python generate_document.py --help

Requirements:
    pip install python-docx
"""

import argparse
import sys
from datetime import datetime

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Error: python-docx is required. Install with: pip install python-docx")
    sys.exit(1)


# Hebrew month names
HEBREW_MONTHS = [
    "ינואר", "פברואר", "מרס", "אפריל", "מאי", "יוני",
    "יולי", "אוגוסט", "ספטמבר", "אוקטובר", "נובמבר", "דצמבר"
]


def format_israeli_date(dt=None):
    """Format date in Israeli convention: DD בMONTH YYYY."""
    if dt is None:
        dt = datetime.now()
    return f"{dt.day} ב{HEBREW_MONTHS[dt.month - 1]} {dt.year}"


def set_rtl_paragraph(paragraph):
    """Set paragraph to RTL direction."""
    pPr = paragraph._p.get_or_add_pPr()
    pPr.set(qn("w:bidi"), "1")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def set_hebrew_font(run, font_name="David", size=12):
    """Set Hebrew font for a text run."""
    run.font.size = Pt(size)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:cs"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def create_formal_letter(output_path):
    """Create a formal Hebrew letter template."""
    doc = Document()

    # Set A4 page size with Israeli margins
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Date
    p = doc.add_paragraph()
    set_rtl_paragraph(p)
    run = p.add_run(format_israeli_date())
    set_hebrew_font(run)

    # Recipient
    p = doc.add_paragraph()
    set_rtl_paragraph(p)
    run = p.add_run("לכבוד")
    set_hebrew_font(run, size=12)

    p = doc.add_paragraph()
    set_rtl_paragraph(p)
    run = p.add_run("[שם הנמען]")
    set_hebrew_font(run, size=12)

    # Subject
    doc.add_paragraph()
    p = doc.add_paragraph()
    set_rtl_paragraph(p)
    run = p.add_run("הנדון: [נושא המכתב]")
    set_hebrew_font(run, size=12)
    run.bold = True

    # Greeting
    doc.add_paragraph()
    p = doc.add_paragraph()
    set_rtl_paragraph(p)
    run = p.add_run("שלום רב,")
    set_hebrew_font(run)

    # Body placeholder
    doc.add_paragraph()
    p = doc.add_paragraph()
    set_rtl_paragraph(p)
    run = p.add_run("[תוכן המכתב]")
    set_hebrew_font(run)

    # Closing
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    set_rtl_paragraph(p)
    run = p.add_run("בכבוד רב,")
    set_hebrew_font(run)

    p = doc.add_paragraph()
    set_rtl_paragraph(p)
    run = p.add_run("[שם החותם]")
    set_hebrew_font(run)

    doc.save(output_path)
    print(f"Created formal letter: {output_path}")


def create_contract(output_path):
    """Create a Hebrew contract template."""
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Title
    p = doc.add_paragraph()
    set_rtl_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("הסכם [סוג ההסכם]")
    set_hebrew_font(run, size=16)
    run.bold = True

    # Date and number
    p = doc.add_paragraph()
    set_rtl_paragraph(p)
    run = p.add_run(f"נערך ונחתם ביום {format_israeli_date()}")
    set_hebrew_font(run)

    # Parties
    doc.add_paragraph()
    p = doc.add_paragraph()
    set_rtl_paragraph(p)
    run = p.add_run("בין:")
    set_hebrew_font(run, size=12)
    run.bold = True

    p = doc.add_paragraph()
    set_rtl_paragraph(p)
    run = p.add_run('[שם צד א\'], ת.ז./ח.פ. [מספר], מרחוב [כתובת] (להלן: "צד א")')
    set_hebrew_font(run)

    p = doc.add_paragraph()
    set_rtl_paragraph(p)
    run = p.add_run("לבין:")
    set_hebrew_font(run, size=12)
    run.bold = True

    p = doc.add_paragraph()
    set_rtl_paragraph(p)
    run = p.add_run('[שם צד ב\'], ת.ז./ח.פ. [מספר], מרחוב [כתובת] (להלן: "צד ב")')
    set_hebrew_font(run)

    # Clauses
    for i in range(1, 4):
        doc.add_paragraph()
        p = doc.add_paragraph()
        set_rtl_paragraph(p)
        run = p.add_run(f"{i}. [סעיף {i}]")
        set_hebrew_font(run)

    # Signatures
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    set_rtl_paragraph(p)
    run = p.add_run("ולראיה באו הצדדים על החתום:")
    set_hebrew_font(run)

    doc.add_paragraph()
    p = doc.add_paragraph()
    set_rtl_paragraph(p)
    run = p.add_run("_______________          _______________")
    set_hebrew_font(run)

    p = doc.add_paragraph()
    set_rtl_paragraph(p)
    run = p.add_run('     צד א\'                          צד ב\'')
    set_hebrew_font(run)

    doc.save(output_path)
    print(f"Created contract template: {output_path}")


def main():
    parser = argparse.ArgumentParser(
        description="Generate Hebrew RTL Word documents"
    )
    parser.add_argument(
        "--template", choices=["letter", "contract"],
        default="letter",
        help="Document template to generate (default: letter)"
    )
    parser.add_argument(
        "--output", default="document.docx",
        help="Output file path (default: document.docx)"
    )
    args = parser.parse_args()

    if args.template == "letter":
        create_formal_letter(args.output)
    elif args.template == "contract":
        create_contract(args.output)


if __name__ == "__main__":
    main()
